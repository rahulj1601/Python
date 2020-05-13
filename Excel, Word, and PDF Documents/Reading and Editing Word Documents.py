import docx

d = docx.Document('/Users/rahul/Documents/Programming/Automating the Boring Stuff with Python /excel, word, and pdf documents/demo.docx')
#creates a new document object

print(d.paragraphs) #list of paragraph objects

print(d.paragraphs[0].text) #refers to text of first paragraph object

p = d.paragraphs[1]

print(p.runs) #number of different style changes in a paragraph

print(p.runs[0].text) #the different groups of text with different runs
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)

p.runs[0].bold #returns true because the first run of text is bold
p.runs[3].italic #returns true or false based on whether the text is italic or not
p.runs[3].underline = True #sets the text for the 3rd run to be underlined

p.runs[3].text = 'italic and underlined.' #alters the text to read the string

d.save('/Users/rahul/Documents/Programming/Automating the Boring Stuff with Python /excel, word, and pdf documents/demo1.docx')

#word documents come with specified styles for certain aspects of the document
p.style = 'Title' #changes the style of the text to a title

d.save('/Users/rahul/Documents/Programming/Automating the Boring Stuff with Python /excel, word, and pdf documents/demo2.docx')

##########################################################################################################################################

d = docx.Document() #create new document
d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('This is another paragraph')

d.save('/Users/rahul/Documents/Programming/Automating the Boring Stuff with Python /excel, word, and pdf documents/demo3.docx')

p = d.paragraphs[0]
p.add_run('This is a new run.')
print(p.runs)
p.runs[1].bold = True
d.save('/Users/rahul/Documents/Programming/Automating the Boring Stuff with Python /excel, word, and pdf documents/demo4.docx')

